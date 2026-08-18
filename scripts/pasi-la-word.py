# -*- coding: utf-8 -*-
"""
Transformă `continut/pasi-proiect.md` într-un fișier Word pentru editare.

DE CE EXISTĂ. Textul celor 40 de pași e ținut în markdown, fiindcă de acolo
ajunge în `js/pasi-proiect.js` și mai departe pe cele trei pagini. Dar Lucian
editează în Word. Scriptul face drumul md → docx ori de câte ori e nevoie, ca
să nu existe două texte care se despart unul de altul.

CUM SE RULEAZĂ (din rădăcina repo-ului):
    python scripts/pasi-la-word.py

Cere `python-docx`, instalat pe 18 august:
    pip install python-docx

⚠️ Drumul invers, din Word înapoi în markdown, NU e automat. După editare,
   fișierul .docx se citește și se trec modificările în `continut/pasi-proiect.md`
   de mână, ca să nu se strecoare formatări ciudate în sursă.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

RADACINA = Path(__file__).resolve().parent.parent
SURSA = RADACINA / "continut" / "pasi-proiect.md"
TINTA = RADACINA / "continut" / "Pasi proiect ApartamenTUal (pentru editare).docx"

CARAMIZIU = RGBColor(0xC2, 0x60, 0x4A)   # accentul platformei
GRI_INCHIS = RGBColor(0x33, 0x33, 0x33)
GRI_MEDIU = RGBColor(0x6A, 0x6A, 0x6A)

FONT_CORP = "Georgia"
FONT_TITLU = "Georgia"


def curata_inline(text):
    """Scoate marcajele markdown care nu se transferă în Word."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def scrie_cu_bold(paragraf, text):
    """Împarte textul pe **bold** și scrie fiecare bucată ca run separat."""
    for bucata in re.split(r"(\*\*[^*]+\*\*)", curata_inline(text)):
        if not bucata:
            continue
        run = paragraf.add_run(bucata.strip("*") if bucata.startswith("**") else bucata)
        run.bold = bucata.startswith("**")


def pregateste_stiluri(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CORP
    normal.font.size = Pt(11)
    normal.font.color.rgb = GRI_INCHIS
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for nivel, marime in ((1, 20), (2, 13), (3, 12)):
        stil = doc.styles[f"Heading {nivel}"]
        stil.font.name = FONT_TITLU
        stil.font.size = Pt(marime)
        stil.font.color.rgb = CARAMIZIU if nivel == 1 else GRI_INCHIS
        stil.font.bold = True
        stil.paragraph_format.space_before = Pt(18 if nivel == 1 else 14)
        stil.paragraph_format.space_after = Pt(6)

    for sectiune in doc.sections:
        sectiune.left_margin = Cm(3)
        sectiune.right_margin = Cm(3)
        sectiune.top_margin = Cm(2.5)
        sectiune.bottom_margin = Cm(2.5)


def construieste(linii, doc):
    in_tabel = False
    tampon = []

    def goleste_tampon():
        """Un paragraf de corp se poate întinde pe mai multe linii în markdown."""
        if not tampon:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        scrie_cu_bold(p, " ".join(tampon))
        tampon.clear()

    for linie_bruta in linii:
        linie = linie_bruta.rstrip()
        gol = not linie.strip()

        # Tabelele din preambul devin linii simple, cu separator vizibil.
        if linie.startswith("|"):
            goleste_tampon()
            if set(linie.replace("|", "").strip()) <= set("-: "):
                continue
            celule = [c.strip() for c in linie.strip("|").split("|")]
            if not in_tabel:
                in_tabel = True
            p = doc.add_paragraph(style="List Bullet")
            scrie_cu_bold(p, "  ·  ".join(c for c in celule if c))
            continue
        in_tabel = False

        if gol:
            goleste_tampon()
            continue

        if linie.startswith("---"):
            goleste_tampon()
            continue

        if linie.startswith("#"):
            goleste_tampon()
            nivel = len(linie) - len(linie.lstrip("#"))
            titlu = curata_inline(linie.lstrip("#").strip())
            doc.add_heading(titlu, level=min(nivel, 3))
            continue

        # „**Sub:** ..." devine linia de context, italică, sub titlul pasului.
        if linie.startswith("**Sub:**"):
            goleste_tampon()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(curata_inline(linie[len("**Sub:**"):].strip()))
            run.italic = True
            run.font.color.rgb = GRI_MEDIU
            continue

        # Linie întreagă în italice: durata fazei.
        if linie.startswith("*") and linie.endswith("*") and not linie.startswith("**"):
            goleste_tampon()
            p = doc.add_paragraph()
            run = p.add_run(curata_inline(linie.strip("*")))
            run.italic = True
            run.font.color.rgb = GRI_MEDIU
            continue

        if re.match(r"^[-*] ", linie):
            goleste_tampon()
            p = doc.add_paragraph(style="List Bullet")
            scrie_cu_bold(p, linie[2:])
            continue

        if re.match(r"^\d+\. ", linie):
            goleste_tampon()
            p = doc.add_paragraph(style="List Number")
            scrie_cu_bold(p, re.sub(r"^\d+\.\s*", "", linie))
            continue

        tampon.append(linie.strip())

    goleste_tampon()


def main():
    if not SURSA.exists():
        sys.exit(f"Nu găsesc sursa: {SURSA}")

    doc = Document()
    pregateste_stiluri(doc)
    construieste(SURSA.read_text(encoding="utf-8").splitlines(), doc)

    TINTA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TINTA)

    # ⚠️ Mesajele de aici sunt fără diacritice dinadins. Consola Windows
    #    rulează pe cp1252 și crapă cu UnicodeEncodeError la primul „ș".
    pasi = SURSA.read_text(encoding="utf-8").count("**Sub:**")
    print(f"Scris: {TINTA}")
    print(f"Pasi gasiti in sursa: {pasi}")


if __name__ == "__main__":
    main()
